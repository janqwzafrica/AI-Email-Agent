import io

from pypdf import PdfReader
from docx import Document


class ExtractionError(Exception):
    """Raised when a document's text can't be extracted."""


def extract_text(file_bytes, filename):
    """
    Extract plain text from a PDF or Word document.

    :param file_bytes: raw bytes of the uploaded file
    :param filename: original filename (used to determine file type)
    :return: extracted text as a string
    """
    ext = filename.rsplit(".", 1)[-1].lower()

    if ext == "pdf":
        return _extract_pdf_text(file_bytes)
    elif ext in ("doc", "docx"):
        return _extract_docx_text(file_bytes, ext)
    else:
        raise ExtractionError(f"Unsupported file extension: {ext}")


def _extract_pdf_text(file_bytes):
    try:
        reader = PdfReader(io.BytesIO(file_bytes))
    except Exception as e:
        raise ExtractionError(f"Could not read PDF: {e}")

    pages_text = []
    for page in reader.pages:
        try:
            pages_text.append(page.extract_text() or "")
        except Exception:
            continue  # skip unreadable page rather than failing the whole doc

    text = "\n".join(pages_text)
    return _normalize_text(text)


def _extract_docx_text(file_bytes, ext):
    if ext == "doc":
        # python-docx only supports .docx (zip-based). Old binary .doc isn't parseable
        # without an external converter (e.g. LibreOffice headless or antiword).
        raise ExtractionError(
            "Legacy .doc files aren't supported yet — please upload as .docx or PDF."
        )

    try:
        document = Document(io.BytesIO(file_bytes))
    except Exception as e:
        raise ExtractionError(f"Could not read Word document: {e}")

    paragraphs = [p.text for p in document.paragraphs]

    # Also pull text out of tables, since some docs put content there
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text)

    text = "\n".join(paragraphs)
    return _normalize_text(text)


def _normalize_text(text):
    if not text:
        return ""
    lines = [line.strip() for line in text.splitlines()]
    lines = [line for line in lines if line]
    normalized = "\n".join(lines)

    if not normalized.strip():
        raise ExtractionError(
            "No readable text found in the document (it may be a scanned image without OCR)."
        )

    return normalized